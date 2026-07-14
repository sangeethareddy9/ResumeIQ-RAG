"""
src/parser.py

Extracts raw text from uploaded resume files (PDF and DOCX).

Design notes:
- Every function here returns a ParsedDocument, never raises. A single
  corrupt file in a batch of 20 resumes should never crash the whole upload --
  the caller just sees status=CORRUPT_OR_UNREADABLE for that one file and
  moves on.
- Text extraction is intentionally "dumb" here (no cleaning) -- cleaning
  happens separately in preprocess.py (Phase 1, next file) so this module
  has one job: get raw text out of a file, reliably.
"""

from __future__ import annotations

import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.models import FileType, ParsedDocument, ParseStatus


def _detect_file_type(filename: str) -> FileType | None:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return FileType.PDF
    if suffix == ".docx":
        return FileType.DOCX
    return None


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a PDF file given as raw bytes."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.PDF,
            status=ParseStatus.CORRUPT_OR_UNREADABLE,
            error_message=f"Could not open PDF: {exc}",
        )

    try:
        pages_text = [page.get_text() for page in doc]
    except Exception as exc:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.PDF,
            status=ParseStatus.ERROR,
            error_message=f"Failed extracting text: {exc}",
        )
    finally:
        doc.close()

    text = "\n".join(pages_text).strip()

    if not text:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.PDF,
            status=ParseStatus.EMPTY,
            error_message="No extractable text found (likely a scanned/image-only PDF).",
        )

    return ParsedDocument(
        filename=filename,
        file_type=FileType.PDF,
        status=ParseStatus.OK,
        raw_text=text,
        char_count=len(text),
    )


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a DOCX file given as raw bytes."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.DOCX,
            status=ParseStatus.CORRUPT_OR_UNREADABLE,
            error_message=f"Could not open DOCX: {exc}",
        )

    try:
        paragraphs = [p.text for p in doc.paragraphs]
        # Also pull text out of tables -- many resumes use tables for layout
        # (e.g. skills columns), and skipping them silently drops content.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
    except Exception as exc:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.DOCX,
            status=ParseStatus.ERROR,
            error_message=f"Failed extracting text: {exc}",
        )

    text = "\n".join(p for p in paragraphs if p.strip()).strip()

    if not text:
        return ParsedDocument(
            filename=filename,
            file_type=FileType.DOCX,
            status=ParseStatus.EMPTY,
            error_message="No extractable text found in document.",
        )

    return ParsedDocument(
        filename=filename,
        file_type=FileType.DOCX,
        status=ParseStatus.OK,
        raw_text=text,
        char_count=len(text),
    )


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """
    Main entry point. Detects file type from filename extension and
    routes to the correct parser. Use this from Streamlit/FastAPI rather
    than calling parse_pdf/parse_docx directly.
    """
    file_type = _detect_file_type(filename)

    if file_type is None:
        return ParsedDocument(
            filename=filename,
            file_type=None,
            status=ParseStatus.UNSUPPORTED_TYPE,
            error_message=f"Unsupported file type for '{filename}'. Only .pdf and .docx are supported.",
        )

    if file_type == FileType.PDF:
        return parse_pdf(file_bytes, filename)

    return parse_docx(file_bytes, filename)